# Seleinium (Browser Crawling) - Dynamic Crawler
# E-mail
import os
import re
import smtplib
import time
from datetime import date, datetime
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

# Docx (Word Document)
import docx
# Chart Data
import FinanceDataReader as fdr
# Standard
import numpy as np
import pandas as pd
import requests
# Crawling - Static Crawler
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.dml import MSO_THEME_COLOR_INDEX  # for hyperlink
from docx.shared import Pt  # for font style and size
from docx.shared import RGBColor  # for font color
from docx.shared import Inches
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait
from webdriver_manager.chrome import ChromeDriverManager

stckListing = fdr.StockListing('KRX')
kospi = '01'
kosdaq = '02'

# =============================================
# Regular Expression
# =============================================
words = [
    ## ================
    ## 트렌드 키워드
    ## ================
    '.*그린\s*뉴딜.*',
    '.*그린.*',
    '.*뉴딜.*',
    '.*언택트.*',
    '.*5G.*',
    '.*진단\s*키트.*',
    '.*치료제.*',
    '.*친환경.*',
    '.*2차\s*전지.*',
    '.*이차\s*전지.*',
    '.*배터리.*',
    '.*전기차.*',
    '.*항공\s*우주.*',
    '.*풍력.*',
    '.*태양광.*',
    '.*반도체.*',
    '.*자율\s*주행.*',
    '.*4차\s*산업.*',
    '.*빅\s*데이터.*',
    '.*빅\s*데이타.*',
    
    
    ## ================
    ## 호재 키워드
    ## ================
    '.*공급\s*계약.*',
    '.*성공.*',
    '.*강세.*',
    '.*초고속.*',
    '.*실적\s*개선.*',
    '.*자사주\s*매입.*',
    '.*흑자\s*전환.*',
    '.*호실적.*',
    '.*수혜.*',
    '.*최선호주.*',
    '.*서프라이즈.*',
    '.*기대.*',
    '.*폭증.*',
    '.*핵심.*',
    '.*개선.*',
    '.*주목.*',
    '.*최고.*',
    '.*집중.*',
    '.*달성.*',
    '.*부각.*']

trend_words = [
    ## ================
    ## 트렌드 키워드
    ## ================
    '.*그린\s*뉴딜.*',
    '.*그린.*',
    '.*뉴딜.*',
    '.*언택트.*',
    '.*5G.*',
    '.*진단\s*키트.*',
    '.*치료제.*',
    '.*친환경.*',
    '.*2차\s*전지.*',
    '.*이차\s*전지.*',
    '.*배터리.*',
    '.*전기차.*',
    '.*항공\s*우주.*',
    '.*풍력.*',
    '.*태양광.*',
    '.*반도체.*',
    '.*자율\s*주행.*',
    '.*4차\s*산업.*',
    '.*빅\s*데이터.*',
    '.*빅\s*데이타.*',
    ]

def removeWon(s):
    a = s.strip()
    result = a[:a.find('원')]

    return result

def removeWonForMCap(s):
    a = s.strip()
    result = a[:a.find('억')]

    return result
    

def removeComma(s):
    a = s.strip()
    result = a.replace(",","")

    return result

def calcSMA (values, window):
	weights = np.repeat(1.0, window)/ window
	smas = np.convolve(values, weights, 'valid')
	return smas

def sendEmail(fileName, to_email):
    email_user = 'josephonsk@gmail.com'     
##    email_send = 'joohyeong1211@gmail.com'
##    email_send = '69ij@naver.com'
    email_send = to_email
##    email_send = 'onyoung@chol.com'
    # email_send = 'hwjiyoon@naver.com'
    today = str(date.today())
    # 제목
    subject = today + ' 분석 결과' 

    msg = MIMEMultipart()
    msg['From'] = email_user
    msg['To'] = email_send
    msg['Subject'] = subject

    # 본문 내용
    body = "온석권"
    msg.attach(MIMEText(body,'plain'))

    ############### ↓ 첨부파일이 없다면 삭제 가능  ↓ ########################
    # 첨부파일 경로/이름 지정하기
    filename = fileName  
    attachment = open(filename,'rb')

    part = MIMEBase('application','octet-stream')
    part.set_payload((attachment).read())
    encoders.encode_base64(part)
    part.add_header('Content-Disposition',"attachment", filename= os.path.basename(filename))
    msg.attach(part)
    ############### ↑ 첨부파일이 없다면 삭제 가능  ↑ ########################

    text = msg.as_string()
    server = smtplib.SMTP_SSL('smtp.gmail.com',465)

    server.login(email_user,email_password)

    server.sendmail(email_user,email_send,text)
    server.quit()

def stock_crawler(code):
    #code = 종목번호
    name = code
    base_url = 'https://finance.naver.com/item/coinfo.nhn?code='+ name + '&target=finsum_more'
    
    browser.get(base_url)
    #frmae구조 안에 필요한 데이터가 있기 때문에 해당 데이터를 수집하기 위해서는 frame구조에 들어가야한다.
    browser.switch_to_frame(browser.find_element_by_id('coinfo_cp'))

    
    #재무제표 "연간" 클릭하기
##    browser.find_elements_by_xpath('//*[@class="schtab"][1]/tbody/tr/td[4]')[0].click()
    #재무제표 "분기" 클릭하기
    browser.find_elements_by_xpath('//*[@id="cns_td22"]')[0].click()
    
    html1 = BeautifulSoup(browser.page_source,'html.parser')
    
    #기업명 뽑기
    title0 = html1.find('head').find('title').text
    print(title0.split('-')[-1])
    
    html22 = html1.find('table',{'class':'gHead01 all-width','summary':'주요재무정보를 제공합니다.'})
    
    #date scrapy
    thead0 = html22.find('thead')
    tr0 = thead0.find_all('tr')[1]
    th0 = tr0.find_all('th')
    
    date = []
    for i in range(len(th0)):
        date.append(''.join(re.findall('[0-9/]',th0[i].text)))
    
    #columns scrapy
    tbody0 = html22.find('tbody')
    tr0 = tbody0.find_all('tr')
    
    col = []
    for i in range(len(tr0)):

        if '\xa0' in tr0[i].find('th').text:
            tx = re.sub('\xa0','',tr0[i].find('th').text)
        else:
            tx = tr0[i].find('th').text

        col.append(tx)
    
    #main text scrapy
    td = []
    for i in range(len(tr0)):
        td0 = tr0[i].find_all('td')
        td1 = []
        for j in range(len(td0)):
            if td0[j].text == '':
                td1.append('0')
            else:
                td1.append(td0[j].text)

        td.append(td1)
    
    td2 = list(map(list,zip(*td)))
    return pd.DataFrame(td2,columns = col,index = date)

def trackOrgBuy(market, dayCode):
    ret = []
    base_url = 'https://finance.naver.com/sise/sise_deal_rank.nhn?sosok='+market+'&investor_gubun=1000'
    
    browser.get(base_url)
    
    browser.switch_to_frame(browser.find_element_by_name('buy'))
    soup = BeautifulSoup(browser.page_source,'html.parser')
    data = soup.find('div').find('div').find('div').find_all('div')[dayCode].find_all('table')[1].find_all('tr')

    try:
        for i in range(2, len(data)):
            tdName = data[i].find_all('td')[0].text.strip()
            
            if tdName == "":
                continue
            tdSum = int(removeComma(data[i].find_all('td')[2].text.strip()))
            selectedStck = stckListing.loc[stckListing['Name']==tdName]
            
            if selectedStck.empty == False:
                stckSymbol = str(selectedStck['Symbol'].values[0])
                dic = {'code':stckSymbol, 'sum':tdSum, 'count':1}
                ret.append(dic)

    except Exception as e:
        print(e)
        
    return ret


def trackForBuy(market, dayCode):
    ret = []
    base_url = 'https://finance.naver.com/sise/sise_deal_rank.nhn?sosok='+market+'&investor_gubun=9000'
    
    browser.get(base_url)
    
    browser.switch_to_frame(browser.find_element_by_name('buy'))
    soup = BeautifulSoup(browser.page_source,'html.parser')
    data = soup.find('div').find('div').find('div').find_all('div')[dayCode].find_all('table')[1].find_all('tr')

    try:
        for i in range(2, len(data)):
            tdName = data[i].find_all('td')[0].text.strip()
            
            if tdName == "":
                continue
            tdSum = int(removeComma(data[i].find_all('td')[2].text.strip()))
            selectedStck = stckListing.loc[stckListing['Name']==tdName]
            
            if selectedStck.empty == False:
                stckSymbol = str(selectedStck['Symbol'].values[0])
                dic = {'code':stckSymbol, 'sum':tdSum, 'count':1}
                ret.append(dic)

    except Exception as e:
        print(e)
        
    return ret
    


def newsFinder(code):
    newsList = []
    base_url = 'https://finance.naver.com/item/news.nhn?code='+code
    browser.get(base_url)

    browser.switch_to_frame(browser.find_element_by_id('news_frame'))
    soup = BeautifulSoup(browser.page_source,'html.parser')
    data = soup.find('div').find('table').find('tbody').find_all('tr')

    for d in data:
        a = d.find('td').find('a').text
        b = 'https://finance.naver.com' + d.find('td').find('a')['href']
        
        if a in newsList:
            continue
        for w in words:
            gNews = re.search(w, a)

            if gNews:
                date = d.find('td',{'class':'date'}).text[:11]
                a += ' '+"["+date+"]"
                dic = {'title':a, 'href':b}
                newsList.append(dic)
                break
            
    return newsList

def mergeBuyList(o1, o2):
    found = False
    for a in o1:
        for b in o2:
            if a['code'] == b['code']:
                b['sum'] += a['sum']
                b['count'] += a['count']
                found = True
                break
        if found:
            found = False
        else:
            o2.append(a)

    return o2

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink
        
    
if __name__ == "__main__":
    
    ## 2번 이상
    ## 코스피 7만이상 코스닥 3.5만이상
    ## 재무재표 통과 (부채비율, 자본유보율)
    ## 차트 통과 (20,10,5 에서 5가 올라가는중?)
    result = ""
    document = Document()
    document.add_heading('R e p o r t',0).alignment = 1 # left:0, center:1, right:2

    df_reverse = fdr.DataReader('KS11')
    df = df_reverse.iloc[::-1]
    today = str(df.head().index.values[0])[:10]
    yesterday = str(df.head().index.values[1])[:10]
                                           
    document.add_paragraph(yesterday + ' ~ ' + today).alignment = 2
    
    browser = webdriver.Chrome(ChromeDriverManager().install())
    browser.maximize_window()
    
    count = 0
    tickers = fdr.StockListing('KRX')['Symbol'].values

    for ticker in tickers:
        try:
            count+=1
            print(count)
            df_reverse = fdr.DataReader(ticker)
            df = df_reverse.iloc[::-1]

            selectedStck = stckListing.loc[stckListing['Symbol']==ticker]
            stckMarket = str(selectedStck['Market'].values[0])

            if stckMarket == 'KONEX':
                continue
            
            stckName = str(selectedStck['Name'].values[0])
            
            base_url = 'https://navercomp.wisereport.co.kr/v2/company/c1010001.aspx?cmp_cd='+str(ticker)+'&cn='
            res = requests.get(base_url)
            res.raise_for_status()
            soup = BeautifulSoup(res.text, 'html.parser')

        
            sma_df = df['Close']
            sma_5 = calcSMA(sma_df, 5)
            sma_10 = calcSMA(sma_df, 10)
            sma_20 = calcSMA(sma_df, 20)
            
            ## =======================================================================
            ##                                 PRICE not correct value
            ## =======================================================================
            ## Stock Price
            price = float(removeComma(removeWon(soup.select("#cTB11 > tbody > tr:nth-child(1) > td")[0].text)))
            
            ## =======================================================================
            ##                                  PER
            ## =======================================================================
            ## EPS
            eps = float(removeComma(soup.select("#pArea > div.wrapper-table > div > table > tr:nth-child(3) > td > dl > dt:nth-child(1) > b")[0].text))
            ## PER
            per = price/eps
            ## Industry Avg PER
##            industryPER = float(removeComma(soup.select("#pArea > div.wrapper-table > div > table > tr:nth-child(3) > td > dl > dt:nth-child(4) > b")[0].text))
            ## =======================================================================
            ##                                  PBR
            ## =======================================================================
            ## BPS
            bps = float(removeComma(soup.select("#pArea > div.wrapper-table > div > table > tr:nth-child(3) > td > dl > dt:nth-child(2) > b")[0].text))
            ## PBR
            pbr = price/bps

            ## =======================================================================
            ##                                 M.Cap
            ## =======================================================================
            ## Market Cap
            marketCap = float(removeComma(removeWonForMCap(soup.select("#cTB11 > tbody > tr:nth-child(5) > td")[0].text)))

            numStock = removeComma(soup.select("#cTB11 > tbody > tr:nth-child(7) > td")[0].text)
            cutInd = numStock.find("주")
            numStock = int(numStock[:cutInd])
            

            shareData = soup.find('table', id='cTB13').find('tbody').find_all("tr") # follow newsfinder for more detail
            shareTotal = 0.0
            shareNames = []
            shareNums = []
            
            for d in shareData:
                a = d.find('td').find('span', {"class": "icon-sprite icon-moreE"})
                if not a:
                    a = d.find('td').find('span', {"class": "cut"})
                a = a.text
                if a == u'\xa0':
                    continue
                a = a.replace(u'\xa0', '')
                b = d.find('td', {'class':'line num'}).text
                b = int(removeComma(b.replace(u'\xa0', '')))
                
                shareNum = round((b/numStock)*100,2)
                shareTotal = shareTotal + shareNum

                shareNames.append(a)
                shareNums.append(str(shareNum))

            if shareTotal > 20.0 and pbr * per < 22.5 and per > 0.0:
                company = stock_crawler(str(ticker))
                crr = float(removeComma(company.iloc[4][24])) # 2020/09 자본유보율
                d = float(removeComma(company.iloc[4][23])) # 2020/09 부채비율
                r1 = float(company.iloc[4][21]) # 2020/09 ROE
                r2 = float(company.iloc[3][21]) # 2020/06 ROE
                r3 = float(company.iloc[2][21]) # 2020/03 ROE
                te = float(removeComma(company.iloc[4][9])) # 2020/09 자본총계
                c = float(removeComma(company.iloc[4][12])) # 2020/09 자본금
                tec = te-c # 2020/09 자본잠식
                
                if tec > 0.0 and crr >= 500.0 and d <= 200.0 and ((r1 >= 10.0 or r2 >= 10.0 or r3 >= 10.0) and r1 > 0.0 and r2 > 0.0 and r3 > 0.0): #d = 100이여야함
                    selectedStck = stckListing.loc[stckListing['Symbol']==str(ticker)]
                    stckName = str(selectedStck['Name'].values[0])
                    comp_info = soup.select('#wrapper > div:nth-child(6) > div.cmp_comment > ul > li')
                    isGood = False
                    for info in comp_info:
                        a = info.text
                        for w in trend_words:
                            gTrends = re.search(w, a)
                            if gTrends:
                                isGood = True
                                break
                        else:
                            continue
                        break
                    if isGood:
                        document.add_paragraph("=============================================================").alignment = 1
                        p = document.add_paragraph()
                        p.alignment = 1
                        p = p.add_run(stckName)
                        p.bold = True
                        p.font.size = Pt(15)

                        document.add_heading('PER: '+ "{:.2f}".format(per) + '\tPBR: ' + "{:.2f}".format(pbr) + '\t부채비율: ' + str(d) +'%\t자본유보율: ' + str(crr) + '%' ,3).alignment = 1
                        document.add_heading('ROE(%)', 3).alignment = 1
                        document.add_heading('[2020/03]: '+ str(r3) + '\t\t[2020/06]: '+ str(r2) + '\t\t[2020/09]: '+ str(r1) ,3).alignment = 1
                        price = int((df['High'][0]+df['Low'][0])/2.0)
                        currency = "{:,}".format(price)
                        document.add_heading('매수 적정가: '+currency+"원",3).alignment = 1
                        if sma_20[0] > sma_10[0] and sma_20[0] > sma_5[0]:
                            p = document.add_paragraph()
                            p.alignment = 1
                            p = p.add_run('*적극 매수 추천')
                            p.font.color.rgb = RGBColor(0xFF, 0x63, 0x47)
                            p.bold = True
                            p.font.name = 'Calibri'
                            p.font.size = Pt(13)
                        document.add_paragraph("=============================================================").alignment = 1

                        p = document.add_paragraph()
                        p = p.add_run('주 주 현 황')
                        p.bold = True
                        p.italic = True
                        p.underline = True
                        p.font.size = Pt(13)

                        for i in range(len(shareNums)):
                            p = document.add_paragraph('', style="List Bullet")
                            p = p.add_run(shareNames[i] + ': ' + shareNums[i] + '%')

                        document.add_paragraph("")
                        
                        p = document.add_paragraph()
                        p = p.add_run('기 업 개 요')
                        p.bold = True
                        p.italic = True
                        p.underline = True
                        p.font.size = Pt(13)

                        for info in comp_info:
                            p = document.add_paragraph('', style="List Bullet")
                            p = p.add_run(info.text)
                            

                        document.add_paragraph("")
                        
                        news = newsFinder(ticker)
                        if len(news) > 0:
                            p = document.add_paragraph()
                            p = p.add_run('관 련 뉴 스')
                            p.bold = True
                            p.italic = True
                            p.underline = True
                            p.font.size = Pt(13)
                            for n in news:
                                p = document.add_paragraph(n['title'], style = 'List Bullet')
                                add_hyperlink(p, "  [링크]", n['href'])
                            document.add_paragraph("")

                
        except Exception as e:
            print(ticker)
            print(e)
            continue

    browser.quit()
    document.add_page_break()
    document.save('report_'+today+'.docx')
    sendEmail('report_'+today+'.docx', 'sklass2345@gmail.com')

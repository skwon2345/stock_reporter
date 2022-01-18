# Seleinium (Browser Crawling) - Dynamic Crawler
# E-mail
import os
import re
import smtplib
import time
from datetime import date, datetime
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from os.path import dirname, join

# Docx (Word Document)
import docx
# Chart Data
import FinanceDataReader as fdr
# Database (Firestore)
import firebase_admin
# Standard
import numpy as np
import pandas as pd
import requests
# Scheduler
from apscheduler.schedulers.blocking import BlockingScheduler
# Crawling - Static Crawler
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.dml import MSO_THEME_COLOR_INDEX  # for hyperlink
from docx.shared import Pt  # for font style and size
from docx.shared import RGBColor  # for font color
from docx.shared import Inches
from dotenv import load_dotenv
from firebase_admin import credentials, firestore, initialize_app, storage
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait
from webdriver_manager.chrome import ChromeDriverManager


def removeWon(s):
    a = s.strip()
    result = a[:a.find('원')]

    return result

def removeWonForMCap(s):
    a = s.strip()
    result = a[:a.find('억')]

    return result
    

def removeComma(s):
    result = s.replace(",","")

    return result

def calcSMA (values, window):
	weights = np.repeat(1.0, window)/ window
	smas = np.convolve(values, weights, 'valid')
	return smas

def sendEmail(fileName, to_email):
    email_user = 'josephonsk@gmail.com'     
    email_send = to_email
    today = str(date.today())
    # 제목
    subject = today + ' 분석 결과' 

    msg = MIMEMultipart()
    msg['From'] = email_user
    msg['To'] = email_send
    msg['Subject'] = subject

    # 본문 내용
    body = "온석권"
    msg.attach(MIMEText(body,'plain'))

    ############### ↓ 첨부파일이 없다면 삭제 가능  ↓ ########################
    # 첨부파일 경로/이름 지정하기
    filename = fileName  
    attachment = open(filename,'rb')

    part = MIMEBase('application','octet-stream')
    part.set_payload((attachment).read())
    encoders.encode_base64(part)
    part.add_header('Content-Disposition',"attachment", filename= os.path.basename(filename))
    msg.attach(part)
    ############### ↑ 첨부파일이 없다면 삭제 가능  ↑ ########################

    text = msg.as_string()
    server = smtplib.SMTP_SSL('smtp.gmail.com',465)

    server.login(email_user,email_password)

    server.sendmail(email_user,email_send,text)
    server.quit()

def stock_crawler(code):
    #code = 종목번호
    name = code
    base_url = 'https://finance.naver.com/item/coinfo.nhn?code='+ name + '&target=finsum_more'
    
    browser.get(base_url)
    #frmae구조 안에 필요한 데이터가 있기 때문에 해당 데이터를 수집하기 위해서는 frame구조에 들어가야한다.
    browser.switch_to.frame(browser.find_element_by_id('coinfo_cp'))

    
    #재무제표 "연간" 클릭하기
##    browser.find_elements_by_xpath('//*[@class="schtab"][1]/tbody/tr/td[4]')[0].click()
    #재무제표 "분기" 클릭하기
    browser.find_elements_by_xpath('//*[@id="cns_td22"]')[0].click()

    
    html1 = BeautifulSoup(browser.page_source,'html.parser')
    
    #기업명 뽑기
    title0 = html1.find('head').find('title').text
    print(title0.split('-')[-1])
    
    html22 = html1.find('table',{'class':'gHead01 all-width','summary':'주요재무정보를 제공합니다.'})
    
    #date scrapy
    thead0 = html22.find('thead')
    tr0 = thead0.find_all('tr')[1]
    th0 = tr0.find_all('th')
    
    date = []
    for i in range(len(th0)):
        date.append(''.join(re.findall('[0-9/]',th0[i].text)))
    
    #columns scrapy
    tbody0 = html22.find('tbody')
    tr0 = tbody0.find_all('tr')
    
    col = []
    for i in range(len(tr0)):

        if '\xa0' in tr0[i].find('th').text:
            tx = re.sub('\xa0','',tr0[i].find('th').text)
        else:
            tx = tr0[i].find('th').text

        col.append(tx)
    
    #main text scrapy
    td = []
    for i in range(len(tr0)):
        td0 = tr0[i].find_all('td')
        td1 = []
        for j in range(len(td0)):
            if td0[j].text == '':
                td1.append('0')
            else:
                td1.append(td0[j].text)

        td.append(td1)
    
    td2 = list(map(list,zip(*td)))
    return pd.DataFrame(td2,columns = col,index = date)

def trackOrgBuy(market, dayCode):
    ret = []
    base_url = 'https://finance.naver.com/sise/sise_deal_rank.nhn?sosok='+market+'&investor_gubun=1000'
    
    browser.get(base_url)
    
    browser.switch_to.frame(browser.find_element_by_name('buy'))
    soup = BeautifulSoup(browser.page_source,'html.parser')
    data = soup.find('div').find('div').find('div').find_all('div')[dayCode].find_all('table')[1].find_all('tr')

    try:
        for i in range(2, len(data)):
            tdName = data[i].find_all('td')[0].text.strip()
            
            if tdName == "":
                continue
            tdSum = int(removeComma(data[i].find_all('td')[2].text.strip()))
            selectedStck = stckListing.loc[stckListing['Name']==tdName]
            
            if selectedStck.empty == False:
                stckSymbol = str(selectedStck['Symbol'].values[0])
                dic = {'code':stckSymbol, 'sum':tdSum, 'count':1}
                ret.append(dic)

    except Exception as e:
        print(e)
        
    return ret


def trackForBuy(market, dayCode):
    ret = []
    base_url = 'https://finance.naver.com/sise/sise_deal_rank.nhn?sosok='+market+'&investor_gubun=9000'
    
    browser.get(base_url)
    
    browser.switch_to.frame(browser.find_element_by_name('buy'))
    soup = BeautifulSoup(browser.page_source,'html.parser')
    data = soup.find('div').find('div').find('div').find_all('div')[dayCode].find_all('table')[1].find_all('tr')

    try:
        for i in range(2, len(data)):
            tdName = data[i].find_all('td')[0].text.strip()
            
            if tdName == "":
                continue
            tdSum = int(removeComma(data[i].find_all('td')[2].text.strip()))
            selectedStck = stckListing.loc[stckListing['Name']==tdName]
            
            if selectedStck.empty == False:
                stckSymbol = str(selectedStck['Symbol'].values[0])
                dic = {'code':stckSymbol, 'sum':tdSum, 'count':1}
                ret.append(dic)

    except Exception as e:
        print(e)
        
    return ret
    

def sortFunction(value):
        return value["sum"]

def sortBuyList(bList):
    sortedBList = sorted(bList, key=sortFunction, reverse= True)

    return sortedBList
    
def newsFinder(code):
    newsList = []
    base_url = 'https://finance.naver.com/item/news.nhn?code='+code
    browser.get(base_url)

    browser.switch_to.frame(browser.find_element_by_id('news_frame'))
    soup = BeautifulSoup(browser.page_source,'html.parser')
    data = soup.find('div').find('table').find('tbody').find_all('tr')

    for d in data:
        a = d.find('td').find('a').text
        b = 'https://finance.naver.com' + d.find('td').find('a')['href']
        
        if a in newsList:
            continue
        for w in words:
            gNews = re.search(w, a)

            if gNews:
                date = d.find('td',{'class':'date'}).text[:11]
                a += ' '+"["+date+"]"
                dic = {'title':a, 'href':b}
                newsList.append(dic)
                break
            
    return newsList

def mergeBuyList(o1, o2):
    found = False
    for a in o1:
        for b in o2:
            if a['code'] == b['code']:
                b['sum'] += a['sum']
                b['count'] += a['count']
                found = True
                break
        if found:
            found = False
        else:
            o2.append(a)

    return o2

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

def uploadToFirebaseStorage(path_to, name, date):
    localFileName = path_to
    blob = bucket.blob('files/'+name)
    blob.upload_from_filename(localFileName)

    blob.make_public()

    print("your file url:", blob.public_url)

    uploadData = {
        'url': blob.public_url,
        'date': date
    }

    ref = db.collection(u'files').document(date)
    ref.set(uploadData)

def mainFunc():
    ## 2번 이상
    ## 코스피 7만이상 코스닥 3.5만이상
    ## 재무재표 통과 (부채비율, 자본유보율)
    ## 차트 통과 (20,10,5 에서 5가 올라가는중?)

    tickers = []

    document = Document()
    document.add_heading('R e p o r t',0).alignment = 1 # left:0, center:1, right:2
    
    df_reverse = fdr.DataReader('KS11')
    df = df_reverse.iloc[::-1]
    today = str(df.head().index.values[0])[:10]
    yesterday = str(df.head().index.values[1])[:10]
                                           
    document.add_paragraph(yesterday + ' ~ ' + today).alignment = 2

    orgKospiYesterday = trackOrgBuy(kospi,0) # 0 is for yesterday
    orgKospiToday = trackOrgBuy(kospi,2) # 2 is for today
    orgKosdaqYesterday = trackOrgBuy(kosdaq, 0)
    orgKosdaqToday = trackOrgBuy(kosdaq, 2)
    mergedOrgKospi = mergeBuyList(orgKospiYesterday, orgKospiToday)
    mergedOrgKosdaq = mergeBuyList(orgKosdaqYesterday, orgKosdaqToday)

    forKospiYesterday = trackForBuy(kospi, 0)
    forKospiToday = trackForBuy(kospi, 2)
    forKosdaqYesterday = trackForBuy(kosdaq, 0)
    forKosdaqToday = trackForBuy(kosdaq, 2)
    mergedForKospi = mergeBuyList(forKospiYesterday, forKospiToday)
    mergedForKosdaq = mergeBuyList(forKosdaqYesterday, forKosdaqToday)

    mergedKospi = mergeBuyList(mergedOrgKospi, mergedForKospi)
    mergedKosdaq = mergeBuyList(mergedOrgKosdaq, mergedForKosdaq)
    
    tickers.extend(mergedKospi)
    tickers.extend(mergedKosdaq)

    tickers = sortBuyList(tickers)
    
    count = 0
    
    for ticker in tickers:
        count+=1
        print(count)
        df_reverse = fdr.DataReader(ticker['code'])
        df = df_reverse.iloc[::-1]

        selectedStck = stckListing.loc[stckListing['Symbol']==ticker['code']]
        stckMarket = str(selectedStck['Market'].values[0])
        stckName = str(selectedStck['Name'].values[0])

        bombP = 70000
        if stckMarket == 'KOSDAQ':
            bombP = 35000

        sma_df = df['Close']
        sma_5 = calcSMA(sma_df, 5)
        sma_10 = calcSMA(sma_df, 10)
        sma_20 = calcSMA(sma_df, 20)
        
        
        base_url = 'https://navercomp.wisereport.co.kr/v2/company/c1010001.aspx?cmp_cd='+str(ticker['code'])+'&cn='
        res = requests.get(base_url)
        res.raise_for_status()
        soup = BeautifulSoup(res.text, 'html.parser')

        try:
            
            ## =======================================================================
            ##                                 PRICE not correct value
            ## =======================================================================
            ## Stock Price
            price = float(removeComma(removeWon(soup.select("#cTB11 > tbody > tr:nth-child(1) > td")[0].text)))
            
            ## =======================================================================
            ##                                  PER
            ## =======================================================================
            ## EPS
            eps = float(removeComma(soup.select("#pArea > div.wrapper-table > div > table > tr:nth-child(3) > td > dl > dt:nth-child(1) > b")[0].text))
            ## PER
            per = price/eps
            ## Industry Avg PER
##            industryPER = float(removeComma(soup.select("#pArea > div.wrapper-table > div > table > tr:nth-child(3) > td > dl > dt:nth-child(4) > b")[0].text))
            ## =======================================================================
            ##                                  PBR
            ## =======================================================================
            ## BPS
            bps = float(removeComma(soup.select("#pArea > div.wrapper-table > div > table > tr:nth-child(3) > td > dl > dt:nth-child(2) > b")[0].text))
            ## PBR
            pbr = price/bps

            ## =======================================================================
            ##                                 M.Cap
            ## =======================================================================
            ## Market Cap
            marketCap = float(removeComma(removeWonForMCap(soup.select("#cTB11 > tbody > tr:nth-child(5) > td")[0].text)))

            numStock = removeComma(soup.select("#cTB11 > tbody > tr:nth-child(7) > td")[0].text)
            cutInd = numStock.find("주")
            numStock = int(numStock[:cutInd])
            

            shareData = soup.find('table', id='cTB13').find('tbody').find_all("tr") # follow newsfinder for more detail
            shareTotal = 0.0
            shareNames = []
            shareNums = []
            
            for d in shareData:
                a = d.find('td').find('span', {"class": "icon-sprite icon-moreE"})
                if not a:
                    a = d.find('td').find('span', {"class": "cut"})
                a = a.text
                if a == u'\xa0':
                    continue
                a = a.replace(u'\xa0', '')
                b = d.find('td', {'class':'line num'}).text
                b = int(removeComma(b.replace(u'\xa0', '')))
                
                shareNum = round((b/numStock)*100,2)
                shareTotal = shareTotal + shareNum

                shareNames.append(a)
                shareNums.append(str(shareNum))
            
            company = stock_crawler(str(ticker['code']))
##            crr = float(removeComma(company["자본유보율"]["2020/06"]))
##            d = float(removeComma(company["부채비율"]["2020/06"]))
##            r1 = float(company["ROE(%)"]["2020/06"])
##            r2 = float(company["ROE(%)"]["2020/03"])
##            r3 = float(company["ROE(%)"]["2019/12"])

            crr = float(removeComma(company.iloc[4][24])) # 2020/09 자본유보율
            d = float(removeComma(company.iloc[4][23])) # 2020/09 부채비율
            r1 = float(company.iloc[4][21]) # 2020/09 ROE
            r2 = float(company.iloc[3][21]) # 2020/06 ROE
            r3 = float(company.iloc[2][21]) # 2020/03 ROE
            te = float(removeComma(company.iloc[4][9])) # 2020/09 자본총계
            c = float(removeComma(company.iloc[4][12])) # 2020/09 자본금
            tec = te-c # 2020/09 자본잠식

##                p1 = float(removeComma(company["영업이익"]["2017/12"]))
##                p2 = float(removeComma(company["영업이익"]["2018/12"]))
##                p3 = float(removeComma(company["영업이익"]["2019/12"]))
##                p4 = float(removeComma(company["영업이익"]["2020/06"]))
            if shareTotal > 20.0 and tec > 0.0 and crr >= 500.0 and d <= 120.0 and ticker['count'] >= 2: #and r1 >= 10.0 and r2 >= 10.0 and r3 >= 10.0: #d = 100이여야함
                recommended = False
                selectedStck = stckListing.loc[stckListing['Symbol']==str(ticker['code'])]
                stckName = str(selectedStck['Name'].values[0])
                comp_info = soup.select('#wrapper > div:nth-child(6) > div.cmp_comment > ul > li')

                buySum = ''
                if ticker['sum'] >=100:
                    convertUnit = "{:.2f}".format(float(ticker['sum'])/100.0)
                    buySum = convertUnit + '억 (원)'
                else:
                    buySum = str(ticker['sum']*100) + '천만 (원)'

                document.add_paragraph("=============================================================").alignment = 1

                p = document.add_paragraph()
                p.alignment = 1
                p = p.add_run(stckName)
                p.bold = True
                p.font.size = Pt(15)

                p = document.add_paragraph()
                p.alignment = 1
                p = p.add_run('기관+외인 순매수 금액:  '+ buySum)
                p.bold = True
                p.italic = True
                p.underline = True
                p.font.size = Pt(12)
                
                document.add_heading('PER: '+ "{:.2f}".format(per) + '\tPBR: ' + "{:.2f}".format(pbr) + '\t부채비율: ' + str(d) +'%\t자본유보율: ' + str(crr) + '%' ,3).alignment = 1
                document.add_heading('ROE(%)', 3).alignment = 1
                document.add_heading('[2020/03]: '+ str(r3) + '\t\t[2020/06]: '+ str(r2) + '\t\t[2020/09]: '+ str(r1) ,3).alignment = 1

                price = int((df['High'][0]+df['Low'][0])/2.0)
                currency = "{:,}".format(price)
                document.add_heading('매수 적정가: '+currency+"원",3).alignment = 1
                
                if ticker['sum'] > bombP and sma_20[0] > sma_10[0] and sma_20[0] > sma_5[0]:
                    recommended = True
                    p = document.add_paragraph()
                    p.alignment = 1
                    p = p.add_run('*적극 매수 추천')
                    p.font.color.rgb = RGBColor(0xFF, 0x63, 0x47)
                    p.bold = True
                    p.font.name = 'Calibri'
                    p.font.size = Pt(13)
         
                document.add_paragraph("=============================================================").alignment = 1

                p = document.add_paragraph()
                p = p.add_run('주 주 현 황')
                p.bold = True
                p.italic = True
                p.underline = True
                p.font.size = Pt(11)

                for i in range(len(shareNums)):
                    p = document.add_paragraph('', style="List Bullet")
                    p = p.add_run(shareNames[i] + ': ' + shareNums[i] + '%')

                document.add_paragraph("")
                
                p = document.add_paragraph()
                p = p.add_run('기 업 개 요')
                p.bold = True
                p.italic = True
                p.underline = True
                p.font.size = Pt(11)
                
                for info in comp_info:
                    p = document.add_paragraph('', style="List Bullet")
                    p = p.add_run(info.text)

                document.add_paragraph("")
                
                news = newsFinder(ticker['code'])
                if len(news) > 0:
                    p = document.add_paragraph()
                    p = p.add_run('관 련 뉴 스')
                    p.bold = True
                    p.italic = True
                    p.underline = True
                    p.font.size = Pt(11)
                    for n in news:
                        p = document.add_paragraph(n['title'], style = 'List Bullet')
                        add_hyperlink(p, "  [링크]", n['href'])
                    document.add_paragraph("")

                uploadData = {
                    'code': ticker['code'],
                    'name': stckName,
                    'date': today,
                    'bPrice': price,
                    'recommended':recommended,
                    'sum(int)': ticker['sum'],
                    'sum':buySum,
                    'count': ticker['count'],
                    'per': per,
                    'pbr': pbr,
                    'debtRatio(%)': d,
                    'reserveRatio(%)': crr,
                    'roe(%)': {
                        u'2020/09': r1,
                        u'2020/06': r2,
                        u'2020/03': r3
                    }
                }
##                db.collection(u'history').add(uploadData) # without document name (automatically create a document name)
                ref = db.collection(u'history').add(uploadData)
##                ref.set(uploadData)
                


                
        except Exception as e:
            print(ticker)
            print(e)
            continue

##    browser.quit()
    document.add_page_break()
    report_name = 'report_'+today+'.docx'
    report_path = 'reports/'+report_name
    document.save(report_path)
    sendEmail(report_path, 'sklass2345@gmail.com')
    sendEmail(report_path, 'hwjiyoon@naver.com')

    uploadToFirebaseStorage(report_path, report_name, today)
    

if __name__ == "__main__":
    dotenv_path = join(dirname(__file__), '.env')
    load_dotenv(dotenv_path)

    firebase_private_id = os.environ.get("FIREBASE_PRIVATE_ID")
    firebase_private_key1 = os.environ.get("FIREBASE_PRIVATE_KEY1")
    firebase_private_key2 = os.environ.get("FIREBASE_PRIVATE_KEY2")
    firebase_private_key3 = os.environ.get("FIREBASE_PRIVATE_KEY3")
    firebase_private_key4 = os.environ.get("FIREBASE_PRIVATE_KEY4")
    firebase_private_key5 = os.environ.get("FIREBASE_PRIVATE_KEY5")
    firebase_project_id = os.environ.get("FIREBASE_PROJECT_ID")
    email_password = os.environ.get("EMAIL_PASSWORD")
    key = {
      "type": "service_account",
      "project_id": firebase_project_id,
      "private_key_id": firebase_private_id,
      "private_key": "-----BEGIN PRIVATE KEY-----\n"+firebase_private_key1+"\n"+firebase_private_key2+"\n"+firebase_private_key3+"\n"+firebase_private_key4+"\n"+firebase_private_key5+"\nftXJTYjpJRNY+E7Z9iB11URP9HimlXqUBrANMl13WvTq3yTBTCQNq3RSOdv4enh4\nuCtuDOD9AgMBAAECggEAChxO9eqoH2k6GEzssn25vFGFVwb+5qvPLPxfGvEmeSmB\nZf4qx9UAJusRjqX/u2lgBNqDMfcBrI9GBgls7odj2+LIEuXImNF2/dCYVq505cSC\nOR5i0NU/kCxwC5jJZ30HEufapMLvjRlCq+wOz/AdVZhlY4LdBUnKjI1mVfNH9b0i\n763MftES+TcyT3mCl3yT/9rF4q7Yt/4JMEn2gSWsFDsupVoS619nooNeddNj7l0+\nP2liC3lf2XzeyW5RKFTSBbcAtrOv7xnNanIfn5g7Htlyc0lqODgdo1TrEz53l1Pu\naM7NqYgdmjRjhnm9HObUIYS3HLiRMgoMvPI97vZd5QKBgQDUdgtOAGXMu9oC1sCR\nH4KFbESMyVcn4zMDOKqj9MzBY7fCGZHcyt2j0+ScKU4xLR8uw4wS1cH8pFR7je9U\n/IhRY4kvo5btHkFuHhVLAUldNZlgW/M6RSS/Zdr9xXyqhk0AFRTGx3YYzps7vIpZ\nlAajgUvYWYTZSytU9Kg5TnPE+wKBgQDDetZiIvGwUMAiCvTmH5o42Tkk1/liSUJZ\nAlPka2sh6UyZwygR/KCPS4l9jI6RapkhhkZWh62okg8yzyY8f9omIxMVHpOC45wl\nJZQjZnoLis+VBhiFUjYczGAU7l1NcGUzpMoaDpz2WobUBBHA4NlhL/QDbSm5HfgH\nj74dtsfgZwKBgQCDFmJbOjuVWsJkxgQChqKNmAOjqgawgapT7n66sqwzFkem5wfs\nZu7hyjdfcszf4qs+u0CNzn3mV64YnbZxhW7GaNeYISNaIjTUsJMKuDx/2jZMEe+8\noDTQMxQvR85nmH2h81/iVecz3yZAcVAbfQ2rjHczgehJ+es+e8XftdwQSQKBgDG0\nUj6x0/KZFcGSN6Mr9fLuiPEtWWP8bLplNR+nkJ4WzDtsvAuWyNymwFcEHlwrVrSH\nh+sIpMdwHB03PumzOKzpzjhYhfmKFnH7lCdML0wwB4hE/Icp2EiKTwQDeAXwxkrT\nY/AsNw0NB6txPXJlueBUaR9V6n3/eMmzawvzJe+RAoGAb9JSyIt+oF9jwJeCPYJ8\ntPEJaLNWvKqGaulc+uEQm/+m7cP5iTraDg/YOG1ZKzU1f09nw3MkKxJwBcpi/3xz\nm/mBWe31cXSgjcJo5X2BMTZcPUrX9VsIPTjsbs5+E3KFGy6zOm7mI3PkghXEe/0w\n+J1ONZbiFbDggUJpmJYsTjU=\n-----END PRIVATE KEY-----\n",
      "client_email": firebase_project_id+"@appspot.gserviceaccount.com",
      "client_id": "106074353173392870934",
      "auth_uri": "https://accounts.google.com/o/oauth2/auth",
      "token_uri": "https://oauth2.googleapis.com/token",
      "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs",
      "client_x509_cert_url": "https://www.googleapis.com/robot/v1/metadata/x509/"+firebase_project_id+"%40appspot.gserviceaccount.com"
    }

    cred = credentials.Certificate(key)
    firebase_admin.initialize_app(cred, {'storageBucket':firebase_project_id+'.appspot.com'})
    db = firestore.client()
    bucket = storage.bucket()
    
    stckListing = fdr.StockListing('KRX')
    kospi = '01'
    kosdaq = '02'


    # =============================================
    # Regular Expression
    # =============================================
    words = [
        ## ================
        ## 트렌드 키워드
        ## ================
        '.*그린\s*뉴딜.*',
        '.*그린.*',
        '.*뉴딜.*',
        '.*언택트.*',
        '.*5G.*',
        '.*진단\s*키트.*',
        '.*치료제.*',
        '.*친환경.*',
        '.*2차 전지.*',
        
        ## ================
        ## 호재 키워드
        ## ================
        '.*공급\s*계약.*',
        '.*성공.*',
        '.*강세.*',
        '.*초고속.*',
        '.*실적\s*개선.*',
        '.*자사주\s*매입.*',
        '.*흑자\s*전환.*',
        '.*호실적.*',
        '.*수혜.*',
        '.*최선호주.*',
        '.*서프라이즈.*',
        '.*기대.*',
        '.*폭증.*',
        '.*핵심.*',
        '.*개선.*',
        '.*주목.*',
        '.*최고.*',
        '.*집중.*',
        '.*달성.*',
        '.*부각.*']
    s = Service('./chromedriver')
    browser = webdriver.Chrome(service=s)
    browser.maximize_window()
    mainFunc()
#     scheduler = BlockingScheduler()
# ##    scheduler.add_job(func=mainFunc, trigger="interval", minutes=20, id="1")

#     scheduler.add_job(func=mainFunc, trigger="cron", day_of_week="mon-fri", hour="18", minute="30", id="1")
#     try:
#         scheduler.start()
#     except KeyboardInterrupt:
#         pass

